from .IngestorInterface import IngestorInterface
from .QuoteModel import QuoteModel
from typing import List
import docx


class DocxIngestor(IngestorInterface):
    allowed_exensions = ['docx']

    @classmethod
    def parse(cls, path: str) -> List[QuoteModel]:
        doc = docx.Document(path)
        quotes = [None] * len(doc.paragraphs)
        for i, line in enumerate(doc.paragraphs):
            if line.text != '':
                data = line.text.split(' - ')
                quotes[i] = QuoteModel(data[0], data[1])
        return quotes